from __future__ import annotations

import hashlib
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path(r"C:\Users\gvadoskr\Yandex.Disk\2025-2026\ПиДИС\4 курс\ДКИП\ЛР1\ЛР1 (Шаблон).docx")
OUTPUT = ROOT / "reports" / "student"
EXPECTED_REFERENCE_SHA256 = "D003E312886862F97A347E4FFCDE1CDEB5C8A0A311FD99E01CACE1F6C467BCEC"
RED = "ED131C"
INK = "1C1C1C"
LIGHT = "F3F3F5"
LINE = "C8C9CE"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(table) -> None:
    repeat_header(table.rows[0])


def format_run(run, size: float = 11, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, size: float = 11, bold: bool = False, color: str = INK, after: float = 4) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in paragraph.runs:
        format_run(run, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    format_run(run, size=16 if level == 1 else 13, bold=True, color=INK)
    if level == 1:
        paragraph.paragraph_format.left_indent = Cm(0.4)
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), RED)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        head = paragraph.add_run(bold_lead)
        format_run(head, bold=True)
        tail = paragraph.add_run(text[len(bold_lead):])
        format_run(tail)
    else:
        run = paragraph.add_run(text)
        format_run(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        marker = paragraph.add_run("• ")
        format_run(marker, bold=True, color=RED)
        format_run(paragraph.add_run(item))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, RED)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        format_run(run, size=9.5, bold=True, color="FFFFFF")
    set_repeat_table_header(table)
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Cm(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if len(table.rows) % 2 == 1:
                set_cell_fill(cell, LIGHT)
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE})
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            format_run(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_answer_box(doc: Document, prompt: str, lines: int = 4) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16.3)
    header = table.cell(0, 0)
    set_cell_fill(header, LIGHT)
    run = header.paragraphs[0].add_run(prompt)
    format_run(run, size=10, bold=True)
    body = table.cell(1, 0)
    set_cell_border(body, top={"val": "single", "sz": "6", "color": LINE}, bottom={"val": "single", "sz": "6", "color": LINE}, left={"val": "single", "sz": "6", "color": LINE}, right={"val": "single", "sz": "6", "color": LINE})
    for index in range(lines):
        paragraph = body.paragraphs[0] if index == 0 else body.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run("________________________________________________________________________________")
        format_run(run, size=9, color="A9A9AE")
    prevent_row_split(table.rows[0])
    prevent_row_split(table.rows[1])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_title_cell(cell, text: str, size: float = 12, bold: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    format_run(run, size=size, bold=bold)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def update_title_page(doc: Document, lab: dict) -> None:
    report_paragraph = doc.paragraphs[5]
    report_paragraph.text = f"\n\nОтчет по лабораторной работе № {lab['number']}"
    report_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in report_paragraph.runs:
        format_run(run, size=14, bold=True)

    title_table = doc.tables[1]
    set_title_cell(title_table.cell(0, 3), f"«{lab['title']}»", size=12.5)
    set_title_cell(title_table.cell(1, 7), f"{lab['id']} · {lab['course']} курс · {lab['semester']} семестр", size=9, bold=False)
    set_title_cell(title_table.cell(3, 5), "МДК.07.01 Управление и автоматизация баз данных", size=12)
    set_title_cell(title_table.cell(4, 7), f"Учебный блок: {lab['block']}", size=8.5, bold=False)

    # В референсе оставлены технические пустые абзацы под незаполненное название темы.
    # Длинные названия лабораторных используют это место, поэтому лишние интервалы убираются.
    for index in [12, 11, 10, 8, 6]:
        delete_paragraph(doc.paragraphs[index])


def append_report(doc: Document, lab: dict) -> None:
    add_heading(doc, "ПАСПОРТ ЛАБОРАТОРНОЙ РАБОТЫ")
    add_table(doc, ["Параметр", "Значение"], [
        ["Идентификатор", lab["id"]],
        ["Вариант", "Укажите номер 01–30 и значения из соответствующей строки матрицы семестра"],
        ["Платформа", lab["platform"]],
        ["Формируемые компетенции", ", ".join(lab["competencies"])],
        ["Технический результат", lab["artifact"]],
    ], [4.8, 11.5])
    add_body(doc, f"Цель. {lab['objective']}", "Цель.")
    add_body(doc, f"Формируемое умение. {lab['skill']}", "Формируемое умение.")

    add_heading(doc, "1. ИСХОДНЫЕ ДАННЫЕ И УСЛОВИЯ")
    add_bullets(doc, lab["inputs"])
    add_heading(doc, "Файлы стартового пакета", level=2)
    add_table(doc, ["№", "Файл", "Назначение / фактическая версия"], [[str(i), Path(path).name, "Зафиксируйте использованную копию и контрольные сведения"] for i, path in enumerate(lab["inputFiles"], 1)], [1.0, 6.1, 9.2])
    add_heading(doc, "Ограничения", level=2)
    add_bullets(doc, lab["constraints"])

    add_heading(doc, "2. ХОД ВЫПОЛНЕНИЯ")
    execution_rows = []
    for index, step in enumerate(lab["steps"], 1):
        execution_rows.append([str(index), step, "Команда / действие:\n\nПолученный вывод:\n", "□ выполнено\n□ проверено"])
    add_table(doc, ["Шаг", "Что выполнено", "Команда, запрос или доказательство", "Статус"], execution_rows, [1.0, 5.4, 7.3, 2.6])

    add_heading(doc, "3. ТЕХНИЧЕСКИЙ РЕЗУЛЬТАТ И ДОКАЗАТЕЛЬСТВА")
    add_table(doc, ["№", "Требуемое доказательство", "Фактическое доказательство / ссылка на рисунок"], [[str(i), evidence, "Вставьте читаемый вывод, фрагмент кода или снимок"] for i, evidence in enumerate(lab["evidence"], 1)], [1.0, 7.0, 8.3])
    add_answer_box(doc, f"Сформулируйте полученный технический результат: {lab['artifact']}", lines=5)

    add_heading(doc, "4. ПРОВЕРКА РЕЗУЛЬТАТА")
    add_table(doc, ["Контроль", "Ожидалось", "Получено и чем подтверждено"], [
        ["Критерий готовности", lab["reminder"]["successCriterion"], ""],
        ["Повторная проверка", "Результат воспроизводится тем же способом", ""],
        ["Типичная ошибка исключена", lab["reminder"]["typicalError"], ""],
    ], [4.0, 6.2, 6.1])

    add_heading(doc, "5. САМОПРОВЕРКА")
    add_table(doc, ["Критерий", "Да", "Доказательство / исправление"], [[item, "□", ""] for item in lab["checklist"]], [10.6, 1.2, 4.5])

    add_heading(doc, "6. ОТВЕТЫ И ВЫВОД")
    for question in lab["reflection"]:
        add_answer_box(doc, question, lines=3)
    add_answer_box(doc, "Вывод: что сделано, что подтверждено и какое решение принято", lines=5)

    add_heading(doc, "СДАЧА В LMS")
    add_table(doc, ["Параметр", "Требование"], [
        ["Формат", lab["lms"]["format"]],
        ["Имя файла", lab["lms"]["recommendedName"]],
        ["Комплектность", lab["lms"]["submission"]],
    ], [4.0, 12.3])
    add_body(doc, "Перед загрузкой удалите личные секреты, убедитесь в читаемости кода и сохраните редактируемый DOCX.")


def generate_one(lab: dict) -> Path:
    target = OUTPUT / f"{lab['id']}_Шаблон_отчёта.docx"
    shutil.copy2(REFERENCE, target)
    doc = Document(target)
    configure_styles(doc)
    update_title_page(doc, lab)
    append_report(doc, lab)
    doc.core_properties.title = f"{lab['id']} — {lab['title']}"
    doc.core_properties.subject = "МДК.07.01 Управление и автоматизация баз данных"
    doc.core_properties.author = "Московский университет «Синергия»"
    doc.save(target)
    return target


def main() -> None:
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    if sha256(REFERENCE) != EXPECTED_REFERENCE_SHA256:
        raise RuntimeError("Контрольная сумма исходного шаблона изменилась")
    labs = json.loads((ROOT / "content" / "labs.json").read_text(encoding="utf-8"))
    OUTPUT.mkdir(parents=True, exist_ok=True)
    generated = [generate_one(lab) for lab in labs]
    if len(generated) != 22:
        raise RuntimeError(f"Ожидалось 22 файла, создано {len(generated)}")
    if sha256(REFERENCE) != EXPECTED_REFERENCE_SHA256:
        raise RuntimeError("Исходный шаблон был изменён во время генерации")
    print(f"Generated {len(generated)} DOCX reports")


if __name__ == "__main__":
    main()
